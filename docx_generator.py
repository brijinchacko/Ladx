"""
LADX - DOCX Document Generator
Converts AI-generated markdown content into professional .docx files.
"""

import os
import re
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("[LADX] python-docx not installed. DOCX generation disabled.")


GENERATED_DIR = os.path.join(os.path.dirname(__file__), "generated_docs")
os.makedirs(GENERATED_DIR, exist_ok=True)


def markdown_to_docx(content: str, title: str, doc_type: str, project_title: str,
                     hardware_info: dict = None) -> str:
    """
    Convert markdown-formatted AI content to a .docx file.
    Returns the filepath of the generated document.
    """
    if not DOCX_AVAILABLE:
        return None

    doc = Document()

    # ---- Styles ----
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x2D, 0x2D, 0x2D)

    # ---- Title page header ----
    heading = doc.add_heading(title, level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0x00, 0x6D, 0x77)

    # Project info
    info_para = doc.add_paragraph()
    info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info_para.add_run(f"Project: {project_title}")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_para.add_run(f"Generated: {datetime.now().strftime('%B %d, %Y')}")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.add_paragraph("")  # spacing

    # Hardware info table if provided
    if hardware_info:
        doc.add_heading("Hardware Configuration", level=1)
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Light Grid Accent 1'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr = table.rows[0].cells
        hdr[0].text = "Parameter"
        hdr[1].text = "Value"
        for key, val in hardware_info.items():
            row = table.add_row().cells
            row[0].text = str(key)
            row[1].text = str(val)
        doc.add_paragraph("")

    # ---- Parse markdown content into docx ----
    lines = content.split('\n')
    in_code_block = False
    code_lines = []
    in_table = False
    table_rows = []

    for line in lines:
        stripped = line.strip()

        # Code block handling
        if stripped.startswith('```'):
            if in_code_block:
                # End code block
                code_text = '\n'.join(code_lines)
                para = doc.add_paragraph()
                para.style = 'Normal'
                run = para.add_run(code_text)
                run.font.name = 'Consolas'
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x1E, 0x1E, 0x1E)
                para_fmt = para.paragraph_format
                para_fmt.left_indent = Inches(0.3)
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
            continue

        if in_code_block:
            code_lines.append(line)
            continue

        # Table handling
        if stripped.startswith('|') and '|' in stripped[1:]:
            # Skip separator rows
            if re.match(r'^[\|\s\-:]+$', stripped):
                continue
            cells = [c.strip() for c in stripped.split('|')[1:-1]]
            if not in_table:
                in_table = True
                table_rows = []
            table_rows.append(cells)
            continue
        elif in_table:
            # Flush table
            _add_table(doc, table_rows)
            in_table = False
            table_rows = []

        # Headings
        if stripped.startswith('# '):
            doc.add_heading(stripped[2:].strip('*'), level=1)
        elif stripped.startswith('## '):
            doc.add_heading(stripped[3:].strip('*'), level=2)
        elif stripped.startswith('### '):
            doc.add_heading(stripped[4:].strip('*'), level=3)
        elif stripped.startswith('#### '):
            doc.add_heading(stripped[5:].strip('*'), level=4)
        elif stripped.startswith('---') or stripped.startswith('***'):
            # Horizontal rule — just add spacing
            doc.add_paragraph("")
        elif stripped == '':
            continue
        else:
            # Regular paragraph with bold/italic support
            para = doc.add_paragraph()
            _add_formatted_text(para, stripped)

    # Flush any remaining table
    if in_table and table_rows:
        _add_table(doc, table_rows)

    # ---- Save ----
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    safe_title = re.sub(r'[^\w\s-]', '', title).strip().replace(' ', '_')[:50]
    filename = f"{safe_title}_{ts}.docx"
    filepath = os.path.join(GENERATED_DIR, filename)
    doc.save(filepath)

    return filepath


def _add_formatted_text(para, text):
    """Parse markdown bold/italic in text and add as runs."""
    # Split by bold markers (**text**)
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = para.add_run(part[2:-2])
            run.bold = True
        else:
            # Check for italic (*text*)
            sub_parts = re.split(r'(\*[^*]+\*)', part)
            for sp in sub_parts:
                if sp.startswith('*') and sp.endswith('*') and len(sp) > 2:
                    run = para.add_run(sp[1:-1])
                    run.italic = True
                else:
                    if sp:
                        para.add_run(sp)


def _add_table(doc, rows):
    """Add a table from parsed markdown rows."""
    if not rows:
        return
    num_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=0, cols=num_cols)
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, row_data in enumerate(rows):
        row = table.add_row()
        for j, cell_text in enumerate(row_data):
            if j < num_cols:
                row.cells[j].text = cell_text.strip()
                # Bold header row
                if i == 0:
                    for paragraph in row.cells[j].paragraphs:
                        for run in paragraph.runs:
                            run.bold = True

    doc.add_paragraph("")  # spacing after table
